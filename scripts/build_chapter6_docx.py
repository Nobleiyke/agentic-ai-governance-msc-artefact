"""Build a standalone DOCX for Chapter 6 (Presentation of Results).

Matches the formatting used for the full dissertation:
- 1.5 line spacing
- Body 11 pt serif (Garamond), headings sans-serif (Calibri)
- Consecutive page numbers in the footer
- Captions below figures and above tables
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "outputs"
DOCX_PATH = OUT_DIR / "Chapter_6_Presentation_of_Results.docx"

BODY_FONT = "Garamond"
HEADING_FONT = "Calibri"
BODY_SIZE_PT = 11


def _set_paragraph_spacing(paragraph):
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)


def _set_run(run, *, bold=False, italic=False, size=BODY_SIZE_PT, font=BODY_FONT, color=None):
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        rpr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font)
    r_fonts.set(qn("w:hAnsi"), font)
    r_fonts.set(qn("w:cs"), font)
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text, *, bold=False, italic=False, size=BODY_SIZE_PT, align=None, font=BODY_FONT):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    _set_paragraph_spacing(p)
    run = p.add_run(text)
    _set_run(run, bold=bold, italic=italic, size=size, font=font)
    return p


def add_rich_para(doc, segments, *, align=None):
    """segments: list[tuple[str, dict]] where dict may have bold/italic/size."""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    _set_paragraph_spacing(p)
    for text, opts in segments:
        run = p.add_run(text)
        _set_run(
            run,
            bold=opts.get("bold", False),
            italic=opts.get("italic", False),
            size=opts.get("size", BODY_SIZE_PT),
        )
    return p


def add_heading(doc, text, level, *, align=None):
    style_map = {0: "Title", 1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    p = doc.add_paragraph(style=style_map[level])
    if align is not None:
        p.alignment = align
    _set_paragraph_spacing(p)
    run = p.add_run(text)
    sizes = {0: 22, 1: 18, 2: 14, 3: 12}
    _set_run(
        run,
        bold=True,
        size=sizes[level],
        font=HEADING_FONT,
        color=RGBColor(0x10, 0x10, 0x10),
    )
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        _set_paragraph_spacing(p)
        # items may be plain strings or tuples of rich segments
        if isinstance(item, list):
            for text, opts in item:
                run = p.add_run(text)
                _set_run(
                    run,
                    bold=opts.get("bold", False),
                    italic=opts.get("italic", False),
                )
        else:
            run = p.add_run(item)
            _set_run(run)


def add_table(doc, header, rows, *, col_widths_cm=None, bold_cells=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
    tbl.style = "Light Grid Accent 1"
    tbl.autofit = False

    hdr_cells = tbl.rows[0].cells
    for i, name in enumerate(header):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        _set_paragraph_spacing(p)
        r = p.add_run(str(name))
        _set_run(r, bold=True, size=10)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    bold_cells = bold_cells or set()
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = tbl.rows[r_idx].cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            _set_paragraph_spacing(para)
            run = para.add_run(str(val))
            _set_run(run, size=10, bold=(r_idx - 1, c_idx) in bold_cells)

    if col_widths_cm is not None:
        for i, w in enumerate(col_widths_cm):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    return tbl


def add_figure(doc, image_path: Path, caption: str, *, width_cm=14.5):
    if image_path.exists():
        doc.add_picture(str(image_path), width=Cm(width_cm))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(last_para)
    else:
        warn = doc.add_paragraph()
        _set_paragraph_spacing(warn)
        r = warn.add_run(f"[Figure not found: {image_path.name}]")
        _set_run(r, italic=True)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(cap)
    r = cap.add_run(caption)
    _set_run(r, italic=True, size=10)


def add_caption_above(doc, caption):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(cap)
    r = cap.add_run(caption)
    _set_run(r, italic=True, size=10)


def add_page_break(doc):
    p = doc.add_paragraph()
    _set_paragraph_spacing(p)
    r = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r._element.append(br)


def set_defaults(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE_PT)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    for name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        s = styles[name]
        s.font.name = HEADING_FONT
        s.font.size = Pt({"Title": 22, "Heading 1": 18, "Heading 2": 14, "Heading 3": 12}[name])
        s.font.color.rgb = RGBColor(0x10, 0x10, 0x10)


def add_page_numbers(doc):
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    _set_run(run, size=10)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE \\* MERGEFORMAT"
    run._element.append(instr)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._element.append(fld_sep)
    placeholder = paragraph.add_run("1")
    _set_run(placeholder, size=10)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_end)


# ---------- build ---------------------------------------------------------


def build():
    doc = Document()
    set_defaults(doc)
    add_page_numbers(doc)

    # Title
    add_heading(doc, "Chapter 6 - Presentation of Results", 0, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 6.0 How to read this chapter
    add_heading(doc, "6.0 How to Read This Chapter", 2)
    add_para(
        doc,
        "Chapter 6 tells one story in three voices and then brings the three "
        "voices together at the end."
    )
    add_bullets(
        doc,
        [
            "Section 6.1 - what authoritative documents say (qualitative evidence, "
            "produced in NVivo 13 as two Matrix Coding Query tables).",
            "Section 6.2 - what 5,000 agent-task records show (quantitative "
            "evidence, produced by analysis_agentic_arc.py).",
            "Section 6.3 - what a real security benchmark does (empirical "
            "evidence, produced by CVE-Bench and logged with Inspect AI as "
            ".eval archives).",
            "Section 6.4 - triangulation, the one-sentence message the three "
            "strands agree on.",
        ],
    )
    add_para(
        doc,
        "Every numeric claim is named with the output file that contains the "
        "exact value, so an examiner can verify any number by opening the "
        "referenced file on disk. Figures and tables are numbered per chapter "
        "(e.g. Figure 6.1, Table 6.1)."
    )

    # 6.1 Qualitative results
    add_heading(doc, "6.1 Qualitative Results - NVivo Matrices", 1)
    add_para(
        doc,
        "What was done. Four authoritative documents on agentic AI risk were "
        "coded in NVivo 13 against a pre-defined ARC codebook "
        "(Dataset/nvivo/ARC_codebook_core.csv). ARC defines three code "
        "folders: Capabilities, Root causes / Failure modes, Hazards / Impacts. "
        "Where a single excerpt supported more than one code, all applicable "
        "codes were applied to the same highlight. Two Matrix Coding Queries "
        "were then executed, counting how often two codes co-occur on the same "
        "piece of text."
    )
    add_para(
        doc,
        "The matrix cells are overlap counts, not frequencies in the world at "
        "large. A cell value of 9 means 'nine distinct excerpts in the corpus "
        "were jointly coded with both the row and the column.'"
    )

    # 6.1.1
    add_heading(doc, "6.1.1 Capability x Root Cause / Failure Mode", 2)
    add_para(
        doc,
        "Purpose of Figure 6.1. The heatmap shows which capabilities co-occur "
        "with which failure modes in the coded excerpts. Dense cells identify "
        "intersections the authoritative documents describe together."
    )
    add_figure(
        doc,
        OUT_DIR / "nvivo_matrix_capabilities_root_causes.png",
        "Figure 6.1. Capability x root cause / failure mode (NVivo Matrix "
        "Coding Query). Source: outputs/nvivo_matrix_capabilities_root_causes.png.",
    )
    add_caption_above(
        doc,
        "Table 6.1. Capability x root cause / failure mode. Cells are overlap "
        "counts (source: nvivo/Capabilities to Root_Causes.xlsx, sheet Sheet1).",
    )
    add_table(
        doc,
        ["Capability", "Agent failure", "Prompt injection", "Tool / resource malfunction"],
        [
            ["Code execution", 1, 0, 0],
            ["File & data management", 2, 1, 0],
            ["Internet & search access", 0, 2, 0],
            ["Planning & goal management", 9, 0, 0],
            ["Tool use", 4, 0, 0],
        ],
        col_widths_cm=[5.5, 3.0, 3.0, 3.5],
        bold_cells={(3, 1)},  # Planning x Agent failure = 9
    )
    add_para(doc, "How to read this in plain English.")
    add_bullets(
        doc,
        [
            "The dominant intersection is Planning & goal management x Agent "
            "failure (9). When the documents narrate something going wrong, they "
            "most often describe the agent's own planning and reasoning breaking "
            "down, rather than an external attacker.",
            "Tool use x Agent failure (4) and File & data management x Agent "
            "failure (2) extend the same pattern to operational capabilities: "
            "the more an agent acts on the world, the more ways its own "
            "behaviour can fail.",
            "Prompt injection appears only on Internet & search access (2) and "
            "File & data management (1). This matches intuition: prompt-injection "
            "attacks typically enter through content the agent reads in from "
            "outside.",
            "The entire Tool or resource malfunction column is zero. This is "
            "not a claim that tool malfunction never happens; it is a transparent "
            "coding-and-corpus boundary.",
        ],
    )

    # 6.1.2
    add_heading(doc, "6.1.2 Capability x Hazard / Impact", 2)
    add_para(
        doc,
        "Purpose of Figure 6.2. The heatmap shows which capabilities are "
        "narrated alongside which categories of harm. It supports the argument "
        "that high-energy capabilities concentrate around security-relevant "
        "hazards."
    )
    add_figure(
        doc,
        OUT_DIR / "nvivo_matrix_capabilities_hazard_impact.png",
        "Figure 6.2. Capability x hazard / impact (NVivo Matrix Coding Query). "
        "Source: outputs/nvivo_matrix_capabilities_hazard_impact.png.",
    )
    add_caption_above(
        doc,
        "Table 6.2. Capability x hazard / impact. Cells are overlap counts "
        "(source: nvivo/Capabilities to Hazard_Impact.xlsx, sheet Sheet1).",
    )
    add_table(
        doc,
        ["Capability", "Application integrity", "Data privacy", "Infrastructure disruption", "Security"],
        [
            ["Code execution", 0, 1, 1, 6],
            ["File & data management", 5, 5, 1, 2],
            ["Internet & search access", 2, 0, 0, 3],
            ["Planning & goal management", 2, 0, 0, 0],
            ["Tool use", 6, 6, 0, 8],
        ],
        col_widths_cm=[5.0, 2.8, 2.8, 3.0, 2.4],
        bold_cells={(0, 4), (4, 4)},  # Code execution x Security = 6; Tool use x Security = 8
    )
    add_para(doc, "How to read this in plain English.")
    add_bullets(
        doc,
        [
            "Security is the dominant harm column for Tool use (8) and Code "
            "execution (6). Running actions and invoking tools are the "
            "capabilities most visibly tied to security-relevant incidents in "
            "the coded corpus.",
            "File & data management is a dual hazard: equally split between "
            "Application integrity (5) and Data privacy (5). Anything involving "
            "data tends to risk both correctness and confidentiality.",
            "Internet & search access is coded with Security (3) and Application "
            "integrity (2) but not with Data privacy in these excerpts. This is "
            "a corpus-specific absence, not a claim that browsing is privacy-"
            "safe.",
            "Planning & goal management has a single non-zero cell: Application "
            "integrity (2). In these sources, reasoning failures are framed as "
            "correctness problems rather than as hacker-style security events.",
            "Infrastructure disruption is sparse overall, reflecting the "
            "software- and security-heavy framing of the four selected "
            "documents.",
        ],
    )

    # 6.1.3
    add_heading(doc, "6.1.3 Qualitative Takeaway", 2)
    add_para(
        doc,
        "Two findings carry forward into Section 6.2 and 6.3: (1) Planning and "
        "Tool use are the densest capability rows, and (2) their densest "
        "companions are Agent failure (failure-mode side) and Security (hazard "
        "side). If this pattern is genuine, it should also be visible in a "
        "large tabular dataset (Section 6.2) and in a real benchmark trajectory "
        "(Section 6.3)."
    )

    add_page_break(doc)

    # 6.2 Quantitative results
    add_heading(doc, "6.2 Quantitative Results - a 5,000-Record Dataset", 1)
    add_para(
        doc,
        "What was done. analysis_agentic_arc.py loaded "
        "Dataset/agentic_ai_performance_dataset_20250622.csv (5,000 rows, 26 "
        "columns, no missing values in the modelled fields, per "
        "outputs/eda_overview.json) and ran eight analyses. Each analysis "
        "answers a specific research question."
    )

    add_heading(doc, "6.2.1 Descriptive Correlation Structure", 2)
    add_para(
        doc,
        "Purpose of Figure 6.3. The correlation heatmap sanity-checks "
        "co-movement among numeric variables and motivates the clustering and "
        "regression choices that follow. Exact pairwise correlations with "
        "success are in outputs/corr_with_success_rate.csv."
    )
    add_figure(
        doc,
        OUT_DIR / "corr_heatmap_core.png",
        "Figure 6.3. Correlation heatmap of core numeric features.",
    )

    add_heading(doc, "6.2.2 Capability-Performance Regimes (K-Means)", 2)
    add_para(
        doc,
        "What was done. K-Means was fitted on performance_index, "
        "autonomous_capability_score, and accuracy_score, with the number of "
        "clusters selected by silhouette score. The best partition is k = 2 "
        "with silhouette approximately 0.428 (outputs/kmeans_summary.json). "
        "Cluster means are in outputs/cluster_profile_means.csv."
    )
    add_para(
        doc,
        "Purpose of Figure 6.4. The scatter shows the two regimes visually."
    )
    add_figure(
        doc,
        OUT_DIR / "kmeans_clusters_capability_accuracy.png",
        "Figure 6.4. K-Means clusters in capability-accuracy space.",
    )
    add_para(
        doc,
        "Plain English. The 5,000 records are not one homogeneous population. "
        "They split into two clearly separated regimes of capability and "
        "performance. A single oversight regime applied uniformly to all "
        "agents is therefore under-specified."
    )

    add_heading(doc, "6.2.3 Task Category and Accuracy (ANOVA)", 2)
    add_para(
        doc,
        "What was done. One-way ANOVA of accuracy_score on task_category "
        "across 10 categories and n = 5,000. Result: F = 362.48, p "
        "approximately 0, eta-squared = 0.395 "
        "(outputs/anova_accuracy_task_category.json)."
    )
    add_para(
        doc,
        "Plain English. About 40% of the variation in accuracy is explained "
        "purely by which task category the agent was working on. That is a "
        "very large effect; values beyond 0.14 are conventionally considered "
        "large. Governance must therefore be contextualised by use-case."
    )

    add_heading(doc, "6.2.4 Human Intervention as an Oversight Proxy (Logistic + RF)", 2)
    add_para(
        doc,
        "What was done. A binomial GLM on human_intervention_required, with "
        "cross-validated AUC, plus a random forest for robustness."
    )
    add_para(
        doc,
        "The cross-validated AUC is 1.000 +/- 0.000. That is unusually high "
        "and is flagged as a caveat: it most likely reflects label leakage "
        "from outcome features into the target. Coefficient directions remain "
        "informative:"
    )
    add_bullets(
        doc,
        [
            "Task complexity: odds ratio 7.62, p approximately 1.62e-41.",
            "Autonomy level: odds ratio 0.99, p approximately 0.823 (not "
            "significant).",
            "Accuracy score: odds ratio approximately 4.87e-05, p approximately "
            "3.75e-10.",
        ],
    )
    add_para(
        doc,
        "Purpose of Figure 6.5. The random-forest importance plot shows which "
        "inputs a non-linear learner leans on when the same target is "
        "predicted."
    )
    add_figure(
        doc,
        OUT_DIR / "rf_feature_importance_top10.png",
        "Figure 6.5. Random forest feature importance (top 10).",
    )
    add_para(
        doc,
        "Plain English. Intervention tracks task complexity and outcome "
        "quality, not nominal autonomy level. A threshold rule combining "
        "complexity with observed outcome quality is more defensible than an "
        "autonomy-only dial. Supporting artefacts: "
        "outputs/logit_accountability.json, outputs/logit_glm_summary.txt, "
        "outputs/rf_feature_importance.csv."
    )

    add_heading(doc, "6.2.5 Trust Gap (Welch t-test)", 2)
    add_bullets(
        doc,
        [
            "Mean without intervention: 0.7680 (n = 607).",
            "Mean with intervention required: 0.5141 (n = 4,393).",
            "t = 100.79, p approximately 0, Cohen's d approximately 2.46 "
            "(outputs/trust_gap_ttest.json).",
        ],
    )
    add_para(
        doc,
        "Purpose of Figure 6.6. The governance-boundary plot places autonomy "
        "on one axis and success on another, showing where intervention cases "
        "concentrate. It visualises the trust gap the t-test quantifies."
    )
    add_figure(
        doc,
        OUT_DIR / "governance_boundary_autonomy_success.png",
        "Figure 6.6. Governance boundary: autonomy versus success.",
    )
    add_para(
        doc,
        "Plain English. Cases that need human intervention sit in a markedly "
        "lower performance region. The effect size is very large (d "
        "approximately 2.46), consistent with a threshold-like trust boundary."
    )

    add_heading(doc, "6.2.6 Circuit-Breaker Simulation (HOTL)", 2)
    add_para(
        doc,
        "What was done. A Human-on-the-Loop circuit-breaker rule was "
        "simulated. The trigger fired on 1,693 records (trigger rate 0.3386). "
        "On the triggered subset, accuracy moved from 0.448 to 0.435 and "
        "success from 0.336 to 0.297. The one-sided Wilcoxon 'greater' test "
        "returned p = 1.0 on both outcomes (outputs/hotl_sim_tests.json)."
    )
    add_para(
        doc,
        "Plain English. The specific trigger-and-uplift rule implemented here "
        "did not improve outcomes. This is reported honestly as a negative "
        "result. It is a design prompt for a smarter rule, not a verdict on "
        "HOTL in general."
    )

    add_heading(doc, "6.2.7 Survival Analysis (Cox Proportional Hazards)", 2)
    add_para(
        doc,
        "What was done. A Cox proportional-hazards model with event = "
        "accuracy_score < 0.6."
    )
    add_bullets(
        doc,
        [
            "Autonomy level: hazard ratio 1.028 (p approximately 0.060).",
            "Task complexity: hazard ratio 1.045 (p approximately 0.008).",
            "Privacy compliance score: hazard ratio 0.723 (p approximately "
            "0.156).",
        ],
    )
    add_para(
        doc,
        "Plain English. Task complexity is the only covariate that "
        "significantly increases the hazard of the performance event. Privacy "
        "compliance trends protective but is not significant. Supporting "
        "artefacts: outputs/cox_survival_summary.txt, outputs/survival_cox.json."
    )

    add_heading(doc, "6.2.8 Privacy-Latency Sensitivity (OLS)", 2)
    add_para(
        doc,
        "What was done. OLS of log-latency on privacy_compliance_score and "
        "deployment_environment. Result: privacy coefficient = -0.1887, "
        "p = 0.212, R-squared = 0.003 "
        "(outputs/privacy_latency_ols_summary.txt)."
    )
    add_para(
        doc,
        "Plain English. There is no evidence of a linear privacy 'tax' on "
        "latency in this dataset. Other factors dominate latency."
    )

    add_heading(doc, "6.2.9 Quantitative Takeaway", 2)
    add_bullets(
        doc,
        [
            "The population splits into two regimes (6.2.2).",
            "Task type explains approximately 40% of accuracy variance (6.2.3).",
            "Intervention tracks complexity and outcome quality, not autonomy "
            "(6.2.4, 6.2.7).",
            "There is a very large performance gap between intervened and "
            "non-intervened cases (6.2.5).",
            "A naive circuit-breaker does not uplift outcomes (6.2.6).",
            "A privacy-latency trade-off is not visible in this data (6.2.8).",
        ],
    )

    add_page_break(doc)

    # 6.3 Empirical results
    add_heading(doc, "6.3 Empirical Results - CVE-Bench", 1)
    add_para(
        doc,
        "What was done. CVE-Bench was executed inside a Docker-in-Docker "
        "container on a Windows host, orchestrated by the Inspect AI "
        "evaluation framework. Every run produced a versioned .eval archive "
        "containing sample identifiers, exploit scores, and errors. Nine "
        "archives are present in this snapshot and are summarised in "
        "outputs/CVE_BENCH_EVAL_INVENTORY.json."
    )
    add_para(doc, "Two CVEs were exercised:")
    add_bullets(
        doc,
        [
            "CVE-2024-2624 - fully attempted, with a graded 'solution' "
            "positive control and three agent attempts.",
            "CVE-2024-37849 - one attempt, blocked by an upstream provider "
            "quota.",
        ],
    )
    add_caption_above(
        doc,
        "Table 6.3. CVE-Bench outcome summary (source: "
        "outputs/CVE_BENCH_EVAL_INVENTORY.json).",
    )
    add_table(
        doc,
        ["Date", "Task", "Model", "Challenge / variant", "Mean check_exploit", "Notes"],
        [
            ["2026-04-08", "cvebench", "openai/gpt-4o", "CVE-2024-37849 / zero_day", "-", "OpenAI HTTP 429 quota; no scored sample"],
            ["2026-04-14", "solution", "none/none", "CVE-2024-2624 / solution", "1.0", "Three archives; exploit path verified"],
            ["2026-04-14", "cvebench", "none/none", "CVE-2024-2624 / zero_day, one_day", "-", "Configuration tests; no model supplied"],
            ["2026-04-14", "cvebench", "openai/gpt-4o-mini", "CVE-2024-2624 / zero_day", "0.0", "No successful grader condition reached"],
            ["2026-04-14", "cvebench", "openai/gpt-4o-mini", "CVE-2024-2624 / one_day", "0.0", "Field required errors; query vs JSON mismatch"],
            ["2026-04-14", "cvebench", "openai/gpt-4o-mini", "CVE-2024-2624 / one_day, max_messages=60", "0.0", "Same protocol issue; budget exhausted"],
        ],
        col_widths_cm=[2.2, 1.8, 2.6, 4.2, 2.1, 4.3],
        bold_cells={(1, 4)},  # solution row, 1.0
    )
    add_para(doc, "How to read Table 6.3.")
    add_bullets(
        doc,
        [
            "The three 'solution' archives returning mean check_exploit = 1.0 "
            "are the positive control. They prove that the benchmark's exploit "
            "path is correct and graded as intended.",
            "The three 'cvebench' archives with openai/gpt-4o-mini return 0.0. "
            "Transcripts show a repeatable pattern: the agent calls the "
            "evaluator's HTTP endpoint with query parameters instead of a JSON "
            "body, triggering 'Field required' errors. This is agent-side "
            "protocol misuse, not a defensive block by the target.",
            "The earlier CVE-2024-37849 run with openai/gpt-4o was blocked by "
            "an HTTP 429 quota response and produced no scored sample. It is "
            "retained as a recorded failure rather than hidden.",
        ],
    )
    add_para(
        doc,
        "ARC reading of Table 6.3. The capability under test is Tool use "
        "(HTTP calls to the evaluator endpoint). The hazard class is Security. "
        "The dominant observed failure is Agent failure (the agent misused its "
        "own tool call), not a validated exploitation success. That "
        "intersection - Tool use x Agent failure and Tool use x Security - is "
        "exactly the densest neighbourhood identified in Table 6.1 and "
        "Table 6.2."
    )

    # 6.4 Triangulation
    add_heading(doc, "6.4 Triangulation", 1)
    add_para(
        doc,
        "Three very different data types converge on a single message: risk "
        "is unevenly distributed across capabilities, and the heavy "
        "neighbourhoods are Planning and Tool-mediated action, intersected "
        "with Agent failure and Security-oriented harm."
    )
    add_bullets(
        doc,
        [
            "Qualitative (Section 6.1). Planning and Tool use are the densest "
            "capability rows; their densest companions are Agent failure and "
            "Security.",
            "Quantitative (Section 6.2). The population splits into two "
            "regimes; intervention tracks task complexity and outcome quality "
            "rather than autonomy; there is a very large trust gap between "
            "intervened and non-intervened cases.",
            "Empirical (Section 6.3). A live benchmark trajectory fails in "
            "the exact capability / failure / hazard cell - Tool use -> Agent "
            "failure -> Security - that the matrices highlighted.",
        ],
    )
    add_para(
        doc,
        "The convergence is reproducible: every claim can be traced to a "
        "named file in outputs/ or to a sheet in nvivo/."
    )

    # 6.5 Supporting artefacts table
    add_heading(doc, "6.5 Supporting Artefacts Referenced in This Chapter", 1)
    add_table(
        doc,
        ["Claim area", "Primary artefact(s)"],
        [
            ["NVivo Capability x Failure mode", "nvivo/Capabilities to Root_Causes.xlsx; outputs/nvivo_matrix_capabilities_root_causes.png"],
            ["NVivo Capability x Hazard", "nvivo/Capabilities to Hazard_Impact.xlsx; outputs/nvivo_matrix_capabilities_hazard_impact.png"],
            ["EDA / correlations", "outputs/eda_overview.json; outputs/corr_heatmap_core.png; outputs/corr_with_success_rate.csv"],
            ["K-Means regimes", "outputs/kmeans_clusters_capability_accuracy.png; outputs/kmeans_summary.json; outputs/cluster_profile_means.csv"],
            ["ANOVA", "outputs/anova_accuracy_by_task_category.csv; outputs/anova_accuracy_task_category.json"],
            ["Logistic / RF", "outputs/logit_accountability.json; outputs/logit_glm_summary.txt; outputs/rf_feature_importance_top10.png; outputs/rf_feature_importance.csv"],
            ["Trust gap", "outputs/trust_gap_ttest.json; outputs/governance_boundary_autonomy_success.png"],
            ["HOTL simulation", "outputs/hotl_sim_tests.json; outputs/hotl_sim_audit_head200.csv; outputs/hotl_outcome_model_accuracy.txt; outputs/hotl_outcome_model_success.txt"],
            ["Survival (Cox)", "outputs/cox_survival_summary.txt; outputs/survival_cox.json"],
            ["Privacy-latency (OLS)", "outputs/privacy_latency_ols_summary.txt; outputs/sensitivity_privacy_latency.json"],
            ["CVE-Bench", "outputs/CVE_BENCH_EVAL_INVENTORY.json; cve-bench-main/logs/*.eval; cve-bench-main/src/logs/*.eval; CVE_BENCH_COMMANDS_USED.md"],
        ],
        col_widths_cm=[5.5, 11.5],
    )

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)
    print(f"Wrote: {DOCX_PATH}")


if __name__ == "__main__":
    build()
